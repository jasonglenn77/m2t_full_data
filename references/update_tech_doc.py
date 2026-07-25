"""
Refresh the M2T technical reference doc for the 2026-07 stakeholder
distribution.

Loads the prior "Model 2 - revised" doc, applies targeted patches to
stale parameter values, updates the GIS source files table for the new
dataflow schema, replaces the stale 2026-04-25 Appendix A with a
pointer to the current exec summary, and appends new appendices
covering the parallel model, consensus reporting, and the GIS dataflow
integration. Saves as a new dated version so the prior doc stays as
history.

Run from project root:
    python references/update_tech_doc.py
"""

import copy
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SOURCE = (
    "references/M2T — Meter-to-Transformer Mapping Model 2 - revised.docx"
)
OUT = (
    "references/M2T — Meter-to-Transformer Mapping Model - "
    "Technical Reference (July 2026).docx"
)


# --- Text patches applied to paragraph runs, preserving styling ---
TEXT_FIXES = [
    ("5,760 points = 60 days", "4,320 points = 45 days"),
    ("5,760-point signatures", "4,320-point signatures"),
    ("60 days × 96 fifteen-minute slots", "45 days × 96 fifteen-minute slots"),
    (
        "CORRELATION_THRESHOLD (currently 0.96)",
        "CORRELATION_THRESHOLD (currently 0.95)",
    ),
    ("60 days (current)", "45 days (current)"),
    (
        "MIN_OVERLAP_POINTS (default 4,000 ≈ 70% data presence)",
        "MIN_OVERLAP_POINTS (default 3,000 ≈ 70% data presence)",
    ),
    ("at least WINDOW_DAYS (currently 60)", "at least WINDOW_DAYS (currently 45)"),
    ("Need at least 60 daily files", "Need at least 45 daily files"),
]


# --- Config table row updates ---
CONFIG_ROW_UPDATES = {
    "WINDOW_DAYS": "45",
    "SIGNATURE_LENGTH": "4,320",
    "CORRELATION_THRESHOLD": "0.95",
    "MIN_OVERLAP_POINTS": "3,000",
}


def fix_paragraph_text(paragraph, fixes):
    """Apply text fixes to a paragraph. Prefer single-run replacement
    (preserves inline styling); fall back to merged rewrite when the
    match spans multiple runs (loses fine-grained inline styling but
    preserves paragraph-level formatting)."""
    for old, new in fixes:
        if old not in paragraph.text:
            continue
        replaced = False
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                replaced = True
                break
        if replaced:
            continue
        new_text = paragraph.text.replace(old, new)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for r in paragraph.runs[1:]:
                r.text = ""


def _set_cell_text(cell, text):
    """Clear a cell and write a single paragraph of text."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)


def update_config_table(table):
    """If the table matches the config parameter shape (Parameter / Value /
    Effect header), update the rows we know are stale. Returns True if
    matched."""
    if len(table.rows) < 2:
        return False
    header = [c.text.strip() for c in table.rows[0].cells]
    if header[:1] != ["Parameter"]:
        return False
    for row in table.rows[1:]:
        if len(row.cells) < 2:
            continue
        key = row.cells[0].text.strip()
        if key in CONFIG_ROW_UPDATES:
            _set_cell_text(row.cells[1], CONFIG_ROW_UPDATES[key])
    return True


def update_gis_source_files_table(table):
    """Rewrite the two data rows of the GIS Source Files table to reflect
    the schema-adaptive loader and the new dataflow. Returns True if
    matched."""
    if len(table.rows) < 3:
        return False
    header = [c.text.strip() for c in table.rows[0].cells]
    if "File Pattern" not in header:
        return False
    sp_row = table.rows[1]
    tx_row = table.rows[2]
    if len(sp_row.cells) >= 3:
        _set_cell_text(sp_row.cells[0], "ServicePoints*.csv")
        _set_cell_text(
            sp_row.cells[1],
            "Every electric meter with its badge and transformer-bank "
            "references. The loader is schema-adaptive: it accepts both "
            "the legacy CSV format and the newer Power BI / Fabric "
            "dataflow export.",
        )
        _set_cell_text(
            sp_row.cells[2],
            "BADGENUMBER, SPID, TRANSFORMERBANKOBJECTID, POINT_X/Y "
            "(or SP_X/Y), CCBADDRESS1, CCBCITY, FEEDERID; "
            "TRANSBANKTAG and d_FEEDERID used when present",
        )
    if len(tx_row.cells) >= 3:
        _set_cell_text(tx_row.cells[0], "Transformers*.csv")
        _set_cell_text(
            tx_row.cells[1],
            "Transformer metadata. Loader prefers the OBJECTID join to "
            "ServicePoints (100% coverage on the current dataflow), "
            "and falls back to the TAG join for legacy files. A "
            "fallback decoder from prior extracts fills in "
            "STRUCTNO / TAG / d_FEEDERID / d_SUBTYPECD for transformers "
            "already known when those columns are absent from the "
            "current file.",
        )
        _set_cell_text(
            tx_row.cells[2],
            "OBJECTID (preferred) or TAG, LID, VAULTCD, FEEDERID, "
            "TOTALKVA, POINT_X/Y (or XFER_X/Y); STRUCTNO, d_FEEDERID, "
            "d_SUBTYPECD used when present",
        )
    return True


def delete_from_appendix_a_onwards(doc):
    """Find the paragraph starting the (stale) Appendix A and remove it
    plus every subsequent body element. Returns True if found."""
    body = doc.element.body
    found_idx = None
    for i, child in enumerate(list(body)):
        # Only text-bearing paragraphs, not tables here — Appendix A
        # begins with a heading paragraph.
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if text.strip().startswith("Appendix A"):
            found_idx = i
            break
    if found_idx is None:
        return False
    # Remove everything from found_idx to end
    for child in list(body)[found_idx:]:
        # Don't remove sectPr (page settings — usually last)
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    return True


def add_heading(doc, text, level=1, color=(0x1F, 0x3A, 0x5F)):
    h = doc.add_heading(text, level=level)
    r, g, b = color
    for run in h.runs:
        run.font.color.rgb = RGBColor(r, g, b)
    return h


def add_para(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = italic
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    """Append bulleted items. Uses the 'List Bullet' style if the source
    doc has it registered; otherwise falls back to plain paragraphs
    prefixed with a bullet glyph so the visual reads the same."""
    style_available = "List Bullet" in [s.name for s in doc.styles]
    for item in items:
        if style_available:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
        else:
            p = doc.add_paragraph()
            run = p.add_run("• " + item)
            p.paragraph_format.left_indent = Inches(0.25)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(3)


def style_table(table, header_row_count=1):
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(10)
                    if ri < header_row_count:
                        r.bold = True


def append_new_appendices(doc):
    """Append fresh Appendix A (pointer to exec doc) and Appendix B
    (parallel model + consensus + GIS dataflow + streaming refactor)."""

    # --- Appendix A: Current Metrics Pointer ---
    add_heading(doc, "Appendix A. Current Run Metrics", level=1)
    add_para(
        doc,
        "Headline run metrics move week-over-week and are maintained "
        "in the stakeholder-facing executive update rather than in this "
        "reference. For the most recent numbers — total clusters, "
        "cluster purity, transformer completeness, flagged/high-"
        "confidence counts, recommendation-type breakdown, badges "
        "missing from GIS, lat/lon discrepancies, and consensus tier "
        "counts — see the current M2T Executive Update document "
        "distributed alongside this reference.",
    )
    add_para(
        doc,
        "The executive doc's Section 2 (Latest Metrics) is the "
        "authoritative source for those numbers on any given run.",
        italic=True,
    )

    # --- Appendix B: Parallel Model, Consensus, GIS Dataflow, Streaming ---
    add_heading(
        doc,
        "Appendix B. Parallel Model, Consensus Reporting, and GIS Dataflow",
        level=1,
    )
    add_para(
        doc,
        "Since the prior revision of this reference, the model has "
        "grown three architecturally significant additions. This "
        "appendix summarizes each; the config-driven parameters and "
        "operational commands for the production model in Sections 1–14 "
        "above still apply unchanged.",
    )

    # B.1 Parallel Model
    add_heading(doc, "B.1 Parallel Model (v2)", level=2)
    add_para(
        doc,
        "A second parallel model runs alongside the production model "
        "and consumes the same daily voltage parquets. Instead of "
        "computing a single Pearson correlation over a fixed rolling "
        "window, the parallel model maintains a per-pair ledger of "
        "daily Pearson values and applies stability gates to those "
        "values before edge formation. This means a pair's inclusion "
        "as a graph edge depends not just on today's correlation being "
        "high, but on the pair's daily correlations being consistently "
        "high, having low variance, and not trending downward.",
    )
    add_para(
        doc,
        "Pipeline modules:",
        bold=True,
    )
    add_bullets(
        doc,
        [
            "pair_accumulator.py — persistent ledger of daily Pearson "
            "values per pair. Stores the last WINDOW_DAYS values plus "
            "cumulative statistics (n, historical mean, first/last day).",
            "daily_correlate.py — reads one daily parquet, computes "
            "per-pair daily Pearson correlations for all candidate "
            "pairs, and updates the ledger via the streaming update "
            "path.",
            "streaming_update.py — merges today's correlations into the "
            "on-disk ledger by streaming through Parquet chunks. Never "
            "loads the full ~12M-pair ledger into memory; peak memory "
            "~500 MB regardless of ledger size.",
            "backfill_accumulator.py — replays a range of daily "
            "parquets through the streaming update path. Supports "
            "--resume, --from-date, and --to-date arguments.",
            "cluster_from_accumulator.py — applies the stability gates "
            "(minimum days, mean, standard deviation, trend), builds "
            "mutual top-K edges, and produces v2 clusters.",
            "run_v2_pipeline.py — orchestrates the v2 side of the "
            "weekly cadence: clustering, error detection, ranking, run "
            "recording, stability reporting, and enrichment.",
        ],
    )
    add_para(
        doc,
        "Configuration for v2 gate thresholds is in config.py under "
        "the V2_* constants (V2_THRESHOLD, V2_STD_MAX, V2_TREND_MIN, "
        "V2_MIN_DAYS). Outputs land in data/outputs_v2/ and follow the "
        "same file naming and column conventions as the production "
        "outputs so downstream tooling works against either.",
    )

    # B.2 Consensus Reporting
    add_heading(doc, "B.2 Consensus Reporting", level=2)
    add_para(
        doc,
        "consensus_report.py compares the production and parallel "
        "models' high-confidence sets and produces "
        "data/outputs/consensus_report.csv with a CONSENSUS_TIER "
        "column. Tiers, in decreasing order of signal strength:",
    )
    add_bullets(
        doc,
        [
            "both_high_confidence — both models independently pass all "
            "high-confidence gates on the same (BADGE, RECOMMENDED_"
            "TRANSFORMER) pair. Strongest possible signal from voltage "
            "data alone.",
            "v1_high_confidence_v2_seen — production model is high-"
            "confidence; parallel model has the same recommendation at "
            "lower confidence. Strong.",
            "v2_high_confidence_v1_seen — parallel model is high-"
            "confidence; production model has the same recommendation at "
            "lower confidence. Strong.",
            "v1_high_confidence_only / v2_high_confidence_only — one "
            "model is high-confidence and the other doesn't flag the "
            "pair. Moderate; investigate why the other model missed it.",
            "both_seen — both models flag the pair at lower-than-high "
            "confidence. Weak but worth tracking across runs.",
            "v1_only / v2_only — flagged by only one model. Weakest "
            "tier; treat as candidates rather than definite findings.",
        ],
    )
    add_para(
        doc,
        "The consensus report is the recommended entry point for "
        "downstream triage. Filter to CONSENSUS_TIER = "
        "both_high_confidence AND RECOMMENDATION_TYPE = "
        "cross_feeder_likely_gis_error for the highest-signal subset.",
    )

    # B.3 GIS Dataflow Integration
    add_heading(doc, "B.3 GIS Dataflow Integration", level=2)
    add_para(
        doc,
        "The model consumes GIS data from a Power BI / Fabric dataflow. "
        "The dataflow's output — ServicePoints and Transformers CSV "
        "exports — is dropped into GIS_mapping/, and the loader picks "
        "the most-recently-modified files automatically. The dataflow "
        "schema differs from the legacy file format in three ways:",
    )
    add_bullets(
        doc,
        [
            "Coordinate columns are named SP_X/SP_Y (ServicePoints) "
            "and XFER_X/XFER_Y (Transformers) rather than POINT_X/"
            "POINT_Y. The loader auto-renames them to canonical "
            "POINT_X/POINT_Y at read time.",
            "The ServicePoints-to-Transformers join uses "
            "TRANSFORMERBANKOBJECTID → OBJECTID (100% match coverage) "
            "instead of the legacy TRANSBANKTAG → TAG join (~97% "
            "coverage). The loader prefers OBJECTID when available.",
            "Several decoded columns (d_FEEDERID, STRUCTNO, TAG, "
            "d_SUBTYPECD) are not present in the current dataflow "
            "export. The loader falls back to a per-key decoder built "
            "from the newest older file in GIS_mapping/ that contains "
            "them, filling in values for infrastructure already known "
            "at that snapshot. New infrastructure created after the "
            "decoder file's snapshot shows blank for those columns.",
        ],
    )
    add_para(
        doc,
        "The intent is that when the dataflow owner adds the missing "
        "decoded columns back, no code change is required — the loader "
        "sees them and stops using the fallback decoder for those keys.",
    )

    # B.4 Downstream Identifier Columns
    add_heading(doc, "B.4 Downstream Identifier Columns", level=2)
    add_para(
        doc,
        "Enriched output files now include three identifier columns "
        "requested by downstream users:",
    )
    add_bullets(
        doc,
        [
            "BADGE_SPID — the meter's SPID from ServicePoints "
            "(per-service-point, badge-level).",
            "CURRENT_TX_LID — the LID of the meter's currently-assigned "
            "transformer (from Transformers, joined via "
            "TRANSFORMERBANKOBJECTID).",
            "RECOMMENDED_TX_LID — the LID of the transformer the model "
            "recommends the meter be reassigned to.",
        ],
    )
    add_para(
        doc,
        "CURRENT_TX_VAULTCD and RECOMMENDED_TX_VAULTCD were already "
        "present in the enriched output and remain unchanged.",
    )


def main():
    if not os.path.exists(SOURCE):
        raise SystemExit(f"Source doc not found: {SOURCE}")

    doc = Document(SOURCE)

    # 1. Fix stale text in every paragraph body-wide.
    for para in doc.paragraphs:
        fix_paragraph_text(para, TEXT_FIXES)

    # Paragraphs inside table cells too.
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_paragraph_text(para, TEXT_FIXES)

    # 2. Update config and GIS source-files tables.
    matched_config = False
    matched_gis = False
    for tbl in doc.tables:
        if not matched_config and update_config_table(tbl):
            matched_config = True
            continue
        if not matched_gis and update_gis_source_files_table(tbl):
            matched_gis = True
            continue
    if not matched_config:
        print("WARNING: config parameter table not found — WINDOW_DAYS/etc. not updated")
    if not matched_gis:
        print("WARNING: GIS source files table not found — schema description not updated")

    # 3. Remove stale Appendix A and everything after it (before sectPr).
    if not delete_from_appendix_a_onwards(doc):
        print("WARNING: Appendix A start not found — stale metrics left in place")

    # 4. Append fresh Appendix A (pointer) and Appendix B (new architecture).
    append_new_appendices(doc)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
