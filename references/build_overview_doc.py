"""
Generate the M2T model overview as a Word document.

A plain-language explanation of what the model does, for stakeholders who
need the big picture rather than the implementation. Deliberately contains
no run statistics, so it does not go stale between weekly runs.

Run from the project root:
    python references/build_overview_doc.py
"""

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "references/M2T_Model_Overview.docx"
FIGURE = "references/_voltage_signature_figure.png"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x4C, 0x56, 0x5F)


# --------------------------------------------------------------- helpers
def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def keep_with_next(paragraph):
    """Stop Word orphaning a heading at the foot of a page."""
    p_pr = paragraph._p.get_or_add_pPr()
    el = OxmlElement("w:keepNext")
    p_pr.append(el)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY
    keep_with_next(h)
    return h


def add_para(doc, text, size=11, italic=False, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_numbered(doc, title, body, n):
    p = doc.add_paragraph()
    r = p.add_run(f"{n}.  {title}")
    r.bold = True
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(8)
    keep_with_next(p)

    b = doc.add_paragraph()
    rb = b.add_run(body)
    rb.font.size = Pt(11)
    b.paragraph_format.left_indent = Inches(0.28)
    b.paragraph_format.space_after = Pt(6)
    return b


def add_table(doc, rows, widths=None, header=True):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, text in enumerate(row):
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            para = cells[ci].paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(10)
            run.bold = header and ri == 0
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            if header and ri == 0:
                shade_cell(cells[ci], "E8EDF2")
        if widths:
            for ci, w in enumerate(widths):
                cells[ci].width = Inches(w)
    # Repeat the header row if a table spans a page break.
    if header and len(table.rows):
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:tblHeader"))
    return table


# ---------------------------------------------------------------- figure
def build_figure(path):
    """Two meters at different voltage levels rising and falling together,
    against a third that moves independently. Drawn rather than plotted so
    the document builds without a charting library."""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print("  PIL not available - building without the figure.")
        return None

    import math

    W, H = 1700, 620
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("arial.ttf", 22)
        small = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        font = small = ImageFont.load_default()

    COPPER = (160, 82, 45)
    SLATE = (60, 90, 115)
    RULE = (200, 206, 211)
    INK = (110, 122, 132)

    seed = [20260902]

    def rnd():
        seed[0] = (seed[0] * 1103515245 + 12345) & 0x7FFFFFFF
        return seed[0] / 0x7FFFFFFF - 0.5

    def shared(t):
        return (math.sin(t * math.pi * 2 - 1.1) * 0.55
                + math.sin(t * math.pi * 4 + 0.6) * 0.28
                + math.sin(t * math.pi * 6 + 2.2) * 0.10)

    def other(t):
        return (math.sin(t * math.pi * 2.6 + 2.4) * 0.42
                + math.sin(t * math.pi * 5.2 - 0.9) * 0.34
                + math.sin(t * math.pi * 1.3 + 0.2) * 0.24)

    x0, x1, N = 250, W - 60, 300
    rows = [
        (120, shared, COPPER, "METER A"),
        (290, shared, COPPER, "METER B"),
        (490, other, SLATE, "METER C"),
    ]

    for base, fn, color, label in rows:
        d.line([(x0, base), (x1, base)], fill=RULE, width=1)
        pts = []
        for i in range(N + 1):
            t = i / N
            x = x0 + t * (x1 - x0)
            y = base - fn(t) * 68 + rnd() * 4
            pts.append((x, y))
        d.line(pts, fill=color, width=4, joint="curve")
        d.text((x0, base - 108), label, fill=INK, font=small)

    # Bracket marking the two meters that track each other
    bx = 210
    d.line([(bx + 16, 66), (bx, 66), (bx, 344), (bx + 16, 344)],
           fill=COPPER, width=3)
    d.text((30, 190), "SAME", fill=COPPER, font=font)
    d.text((30, 218), "TRANSFORMER", fill=COPPER, font=font)

    d.text((x0, H - 60), "00:00", fill=INK, font=small)
    d.text((x1 - 70, H - 60), "24:00", fill=INK, font=small)

    img.save(path, dpi=(220, 220))
    return path


# ----------------------------------------------------------------- build
def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---- Title -------------------------------------------------------
    title = doc.add_paragraph()
    tr = title.add_run("How the M2T Model Works")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = NAVY
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sr = sub.add_run(
        "Using the voltage signatures of the meters, the model correlates and "
        "clusters them, then compares the result to the GIS mappings."
    )
    sr.font.size = Pt(13)
    sr.font.color.rgb = GREY
    sub.paragraph_format.space_after = Pt(16)

    # ---- Problem -----------------------------------------------------
    add_heading(doc, "The problem it addresses", 1)
    add_para(doc,
        "GIS holds a record of which secondary distribution transformer serves "
        "each electric meter. Confirming that record for a given meter means "
        "sending someone to look at it. There are roughly 218,000 meters, so "
        "confirming all of them that way is not practical.")
    add_para(doc,
        "The goal is full reconciliation: every meter's assignment confirmed, "
        "so that when an outage occurs on a transformer or a feeder, the "
        "affected meters are known. The model's role is to determine where "
        "verification should be directed first.")

    # ---- Idea --------------------------------------------------------
    add_heading(doc, "The idea behind it", 1)
    add_para(doc,
        "Meters served by the same transformer sit behind the same piece of "
        "equipment. When load rises on that transformer, the voltage every "
        "meter behind it sees dips together. When load falls, it recovers "
        "together. Their voltage readings move as a group.")
    add_para(doc,
        "The model reads those readings and looks for meters whose voltage "
        "patterns track one another. It compares the shape of the curve, not "
        "its height, so two meters can sit at different voltage levels and "
        "still be recognized as moving in step. That shared shape is what "
        "indicates a shared transformer.")

    fig = build_figure(FIGURE)
    if fig:
        doc.add_picture(fig, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = add_para(doc,
            "Meters A and B sit at different voltage levels but rise and fall "
            "together. Meter C is nearby but moves independently, so the model "
            "does not group it with them.",
            size=9, italic=True, space_after=14)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Steps -------------------------------------------------------
    add_heading(doc, "What the model does, step by step", 1)
    steps = [
        ("Read the voltage",
         "Every meter already reports voltage at fifteen-minute intervals. The "
         "model uses a rolling window of recent weeks, so it reflects the "
         "system as it is now."),
        ("Narrow to plausible neighbors",
         "A meter can only share a transformer with one that is physically "
         "close, so the model only compares meters within a short distance of "
         "each other. This is a shortlist, not evidence; distance does not "
         "enter the correlation."),
        ("Compare the patterns",
         "For every candidate pair, it measures how closely the two voltage "
         "patterns move together, and keeps only relationships that are strong "
         "and sustained."),
        ("Require agreement in both directions",
         "A pair only survives if each meter counts the other among its own "
         "closest matches. One-sided resemblance is discarded. This is the "
         "strongest filter against false groupings."),
        ("Form groups, then compare against GIS",
         "Surviving relationships link meters into groups. Only at this point "
         "does GIS enter: the model checks whether everyone in a group is "
         "recorded on the same transformer. Where a meter's record disagrees "
         "with its group, that disagreement is the finding."),
    ]
    for i, (t, b) in enumerate(steps, 1):
        add_numbered(doc, t, b, i)
    add_para(doc,
        "The order matters. The model does not read GIS while forming its "
        "groups, so it cannot reproduce what GIS already says. It reaches an "
        "independent result and then compares.", space_after=12)

    # ---- Two models --------------------------------------------------
    add_heading(doc, "Two models, not one", 1)
    add_para(doc,
        "Two versions run in parallel on the same data using different "
        "methods. A finding both versions produce carries more weight than one "
        "either produces alone.")
    add_table(doc, [
        ["Version", "How it works", "What it contributes"],
        ["Version 1",
         "Treats each meter's whole window as one continuous measurement and "
         "asks how closely two meters track overall.",
         "Finds strong relationships across the window."],
        ["Version 2",
         "Measures each day separately, then examines the whole run of daily "
         "results for that pair.",
         "Confirms the relationship is consistently strong, not just strong on "
         "average. A single overall figure can hide a pair that matched well "
         "for part of the window and poorly for the rest."],
    ], widths=[0.9, 2.5, 3.0])
    doc.add_paragraph()

    # ---- Weekly ------------------------------------------------------
    add_heading(doc, "Why it runs every week", 1)
    add_para(doc,
        "A single run is a snapshot. Conditions during one window, or gaps in "
        "the data, can make two unrelated meters resemble each other for a "
        "time.")
    add_para(doc,
        "Running weekly turns that into the strongest filter available. The "
        "model keeps a record of every past run, and a finding only reaches "
        "the verification list if the same grouping and the same conclusion "
        "have held across most of the recent weeks. Persistence is the "
        "evidence. A meter flagged once may be noise; one flagged week after "
        "week, by two independent methods, is worth investigating.", space_after=12)

    # ---- Limits ------------------------------------------------------
    add_heading(doc, "What it can and cannot tell you", 1)
    add_para(doc, "Being clear about the limits is what makes the findings usable.")
    add_table(doc, [
        ["Reliable", "Beyond its reach"],
        ["Meters recorded on the wrong feeder. Voltage patterns should not "
         "track across separate feeders, so when they do, the record is the "
         "likely explanation.\n\n"
         "Meters with no transformer recorded. If an unmapped meter groups "
         "cleanly with mapped ones, that is a gap with a proposed answer.\n\n"
         "Meters missing from GIS entirely, and meters whose recorded location "
         "is far from where their data places them.",
         "Neighboring transformers on the same feeder. These share an "
         "upstream source, so meters behind them can look nearly identical. "
         "The model still detects that the record looks wrong; what it cannot "
         "do is say which transformer is right. Field validation has confirmed "
         "this limitation in practice.\n\n"
         "Meters with sparse or intermittent data. Without enough readings "
         "there is no pattern to compare.\n\n"
         "Fragmented groups. The model often splits one transformer's meters "
         "across several groups. That is a limitation in recognizing them as "
         "one set, not a claim that the records are wrong."],
    ], widths=[3.2, 3.2])
    doc.add_paragraph()

    # ---- Outputs -----------------------------------------------------
    add_heading(doc, "What comes out of it", 1)
    add_para(doc,
        "Each week produces two lists. The verification list holds the meters "
        "with enough evidence behind them to justify a visit now. The full "
        "difference list holds every meter either version disagrees with GIS "
        "on, ranked by strength of evidence, so nothing is hidden and the "
        "remainder can be worked down over time.")
    add_para(doc,
        "Every meter on the verification list is in scope for a visit. The "
        "priority indicates how much the model could determine on its own, and "
        "therefore what the crew should expect to find.")
    add_table(doc, [
        ["Priority", "Category", "What it means"],
        ["1", "Wrong feeder",
         "The meter's behavior matches a transformer on a different feeder "
         "than its record shows. The strongest case the model can make. "
         "Verify, then re-map."],
        ["2", "Unmapped meter",
         "No transformer on record, and both versions agree which one it "
         "belongs to. Verify, then assign."],
        ["3", "Same feeder",
         "The largest group. Both versions agree the record looks wrong, but "
         "the transformers involved share a feeder, so voltage cannot say "
         "which is correct. The visit decides. Ranked by how many weeks the "
         "same finding has repeated."],
        ["4", "Missing feeder data",
         "Held back because feeder information is absent in GIS for one of the "
         "transformers involved. Populating it moves these into a real "
         "category."],
    ], widths=[0.7, 1.5, 4.2])
    doc.add_paragraph()
    add_para(doc,
        "The list is a work queue, not a set of changes to apply directly. "
        "Crews confirm before anything is re-mapped.", space_after=12)

    # ---- After delivery ----------------------------------------------
    add_heading(doc, "What happens after the list goes out", 1)
    add_para(doc,
        "The findings come back, and the model takes account of them. This is "
        "what turns a weekly report into progress rather than the same list "
        "arriving repeatedly.")
    loop = [
        ("The list is loaded and worked",
         "Missing meters and coordinate discrepancies are handled in the "
         "office. The rest are dispatched for a site visit, working down the "
         "priority order."),
        ("The crew confirms what is actually there",
         "They record the transformer the meter is really connected to. That "
         "single value settles it, whichever way it goes."),
        ("If the model was right, GIS is corrected",
         "The model reads a fresh GIS extract every Monday, sees that the "
         "assignment now matches what it recommended, and records the meter as "
         "resolved."),
        ("If GIS was right, the meter is retired from the list",
         "The confirmation is returned to the model, which stops reporting "
         "that meter. A confirmed record is as valuable as a corrected one: "
         "both remove a meter from the unverified population."),
    ]
    for i, (t, b) in enumerate(loop, 1):
        add_numbered(doc, t, b, i)
    add_para(doc,
        "Every verified meter is counted, so the share of the population that "
        "has been settled only increases. That figure, rather than the weekly "
        "count of findings, is the measure of progress.", space_after=14)

    # ---- Closing -----------------------------------------------------
    add_heading(doc, "How the model relates to GIS", 1)
    add_para(doc,
        "The model is built to stand beside GIS, not underneath it. It does "
        "not use GIS to reach its conclusions, which makes the two independent "
        "sources. When they disagree, that disagreement is the useful part, "
        "and either one can turn out to be the one that is wrong. The field "
        "visit settles it, and every settled meter is one more the map can be "
        "relied on for.")

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    if fig and os.path.exists(fig):
        os.remove(fig)


if __name__ == "__main__":
    main()
