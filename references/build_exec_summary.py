"""
Generate an executive-level Word doc summarizing the M2T model's current
state, latest metrics, and how it handles feedback. Designed for
stakeholders / leadership who want the holistic view, not the
implementation details.

Run from project root:
    python references/build_exec_summary.py
"""

import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "references/M2T_Executive_Update_July2026.docx"


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style_table(table, header_row_count=1):
    table.style = "Light Grid Accent 1"
    table.autofit = False
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(10)
                    if ri < header_row_count:
                        r.bold = True


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(3)


def main():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title block ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("M2T Meter-to-Transformer Mapping Model")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run("Executive Update — Late July 2026")
    run.italic = True
    run.font.size = Pt(13)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(
        "Prepared by Jason Glenn | Engineering & Data Team | Confidential"
    ).font.size = Pt(10)
    doc.add_paragraph()

    # ---- 1. Executive Summary ----
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(
        doc,
        "The Meter-to-Transformer (M2T) model identifies which secondary "
        "distribution transformer each single-phase electric meter is "
        "actually connected to, by correlating voltage patterns observed "
        "over time. Its purpose is to validate and improve the GIS-side "
        "meter-to-transformer assignments — surfacing cases where a "
        "meter's voltage behavior tells a different story than what GIS "
        "currently records.",
    )
    add_para(
        doc,
        "Since the last stakeholder update, the model has continued "
        "evolving on two parallel tracks: a refined production model with "
        "a weekly run cadence and a maturing stability ledger, and a "
        "second parallel model used to cross-check the production "
        "model's recommendations. We now have nine weekly runs of "
        "stability data, an active classification system that "
        "distinguishes high-signal from ambiguous recommendations, and "
        "a structured way for field validation results to feed back "
        "into the model. Additionally, the model now consumes a new "
        "GIS dataflow directly and delivers three additional identifier "
        "columns (SPID, TransformerLID, VaultCD) requested by downstream "
        "users.",
    )

    # ---- 2. Latest Metrics ----
    add_heading(doc, "2. Latest Metrics — Run Ending 2026-05-09", level=1)
    add_para(
        doc,
        "The numbers below come from the most recent weekly cycle. "
        "Deltas are shown against the prior stakeholder-facing run "
        "(2026-05-02) to make week-over-week movement visible. The "
        "model's rolling window remains at 45 days at a 0.95 "
        "correlation threshold — the configuration that struck the "
        "best purity/completeness balance in prior tuning.",
    )

    metrics = [
        ["Metric", "Value", "Change vs 2026-05-02"],
        ["Total meters in model", "218,070", "+65"],
        ["Total clusters identified", "52,077", "+496"],
        ["Cluster purity (no false merges)", "91.8%", "unchanged"],
        ["Transformer completeness (no false splits)", "46.7%", "-0.2 pp"],
        ["Total flagged corrections", "11,407", "+46"],
        ["High-confidence corrections (acted on by field)", "1,850", "-9"],
        [
            "Strong-signal corrections (high gap; one-run filter)",
            "2,348",
            "-10",
        ],
        ["Median confidence gap", "0.167", "unchanged"],
        ["75th-percentile confidence gap", "0.400", "unchanged"],
        ["Badges missing from GIS (push to GIS team)", "73", "+4"],
        [
            "Lat/long discrepancies vs GIS (data quality candidates)",
            "1,290",
            "+2",
        ],
        ["Weekly runs in stability ledger", "9", "+2"],
    ]
    table = doc.add_table(rows=len(metrics), cols=3)
    for ri, row in enumerate(metrics):
        for ci, val in enumerate(row):
            table.cell(ri, ci).text = val
    style_table(table)
    table.columns[0].width = Inches(3.3)
    table.columns[1].width = Inches(1.4)
    table.columns[2].width = Inches(1.4)

    add_para(
        doc,
        "Interpretation. Cluster purity tells us how often the model's "
        "groups contain meters from only one transformer (high is good — "
        "no false merges). Transformer completeness tells us how often a "
        "transformer's entire set of meters lands in a single cluster "
        "(high is good — no false splits). The two metrics naturally pull "
        "against each other, and the run-to-run trajectory shows the "
        "model is making sensible trade-offs as we tune the configuration.",
        italic=True,
    )

    # ---- 3. Recommendation Breakdown ----
    add_heading(doc, "3. What the Field Team Sees", level=1)
    add_para(
        doc,
        "Of the 11,407 corrections the model produced this run, only "
        "1,850 pass the high-confidence gates — these are the cases "
        "with strong cluster votes and consistent recommendations across "
        "weekly runs. Within those, the model labels each correction by "
        "the type of evidence it has, which directs the field team to "
        "the most actionable items first.",
    )
    breakdown = [
        ["Recommendation Type", "Total", "High-Confidence", "Field Action"],
        [
            "cross_feeder_likely_gis_error",
            "193",
            "51",
            "Top priority — likely GIS feeder/transformer error",
        ],
        [
            "new_assignment",
            "232",
            "46",
            "Easy wins — first-time mapping for unmapped meters",
        ],
        [
            "unknown_feeder",
            "5",
            "1",
            "Needs GIS data to populate feeder before action",
        ],
        [
            "same_feeder_ambiguous",
            "10,977",
            "1,752",
            "Lower signal — field validation required before action",
        ],
    ]
    table = doc.add_table(rows=len(breakdown), cols=4)
    for ri, row in enumerate(breakdown):
        for ci, val in enumerate(row):
            table.cell(ri, ci).text = val
    style_table(table)
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(0.7)
    table.columns[2].width = Inches(1.1)
    table.columns[3].width = Inches(2.7)

    add_para(
        doc,
        "The cross-feeder and new-assignment categories — 97 high-"
        "confidence cases this week — are the model's strongest "
        "contributions. Cross-feeder recommendations call out cases "
        "where GIS records a meter on the wrong electrical feeder "
        "entirely; new-assignment recommendations are first-time "
        "mappings for meters GIS has no record of yet. Both are "
        "well-suited to direct field or GIS-team action.",
    )

    # ---- 4. How the Model Handles Feedback ----
    add_heading(doc, "4. How the Model Handles Feedback", level=1)
    add_para(
        doc,
        "An important focus over the past several weeks has been "
        "building structured feedback loops between field operations, "
        "GIS, and the model. These mechanisms keep the model honest, "
        "let real-world findings improve future runs, and surface "
        "patterns that voltage data alone cannot resolve.",
    )

    add_heading(doc, "4.1 Field Validation Findings", level=2)
    add_para(
        doc,
        "When field operations investigates a recommendation, the result "
        "becomes useful information regardless of outcome. A confirmed "
        "correction validates the model's signal; a confirmed false "
        "positive points to a category the model struggles with. A "
        "recent example: field operations validated badge 1020669, "
        "where the model had consistently recommended a different "
        "transformer than GIS records. The field confirmed GIS was "
        "correct. The model had labeled this case 'same_feeder_ambiguous,' "
        "which is precisely the category where voltage signatures alone "
        "cannot reliably distinguish adjacent transformers on the same "
        "feeder. The classification working as intended on a real case "
        "is encouraging — it means the model knows where it is reliable "
        "and where it is not.",
    )

    add_heading(doc, "4.2 Suppression List", level=2)
    add_para(
        doc,
        "A user-maintained list lets us suppress specific recommendations "
        "we have decided to set aside — for example, cases that field "
        "operations has already investigated and confirmed do not need "
        "further action. The model still produces the underlying "
        "analysis, but the suppressed entries are kept out of the active "
        "work list and logged separately for audit. If the model later "
        "changes its mind and points to a different transformer for the "
        "same meter, that new recommendation still surfaces — only the "
        "specific suggestion we suppressed is hidden.",
    )

    add_heading(doc, "4.3 Stability Tracking Across Runs", level=2)
    add_para(
        doc,
        "Every weekly run is recorded in a persistent ledger that scores "
        "how consistent the model is over time. After nine runs of "
        "history, recommendations that have appeared in every run get "
        "elevated, and recommendations that have only flickered in once "
        "or twice get downgraded. This means the model genuinely gets "
        "smarter about what to trust as more data accumulates — a "
        "one-off flag is treated as suspect, while a recommendation "
        "that has survived nine weeks of changing data is treated as "
        "robust evidence.",
    )

    add_heading(doc, "4.4 GIS Data Refresh", level=2)
    add_para(
        doc,
        "The model now consumes GIS data directly from a shared Power "
        "BI / Fabric dataflow. When the dataflow refreshes, the model "
        "picks up the newest ServicePoints and Transformers extracts "
        "on the next run — no manual conversion or reformatting is "
        "required. This keeps the model's comparison-to-GIS truth file "
        "current with whatever corrections the GIS team has already "
        "made and delivers 42% more service-point-to-transformer "
        "relationships than the previous file-based approach.",
    )

    add_heading(doc, "4.5 Parallel Model (Consensus Cross-Check)", level=2)
    add_para(
        doc,
        "A second version of the model now runs alongside the production "
        "model. It uses a different mathematical approach to the same "
        "underlying voltage data, with built-in checks for whether each "
        "meter-pair relationship is consistent over time or whether it "
        "looks like coincidence. When both models agree on a "
        "recommendation at high confidence, that becomes the strongest "
        "signal in the entire output — two methodologically-independent "
        "analyses converging on the same answer is much rarer, and much "
        "more reliable, than either alone. When they disagree, the "
        "disagreement itself is a flag worth investigating.",
    )
    add_para(
        doc,
        "The most recent consensus comparison identified 1,739 "
        "recommendations where both models agree at high confidence, "
        "including 47 strong cross-feeder findings and 44 first-time "
        "assignment suggestions. These are the highest-priority items "
        "for field and GIS-team action.",
    )

    # ---- 5. What's New Since the Last Update ----
    add_heading(doc, "5. What's Changed Since the Last Update", level=1)
    add_bullets(
        doc,
        [
            "Integrated the new GIS Power BI / Fabric dataflow. The "
            "model now reads the freshest ServicePoints and Transformers "
            "extracts directly from the dataflow output and joins "
            "service points to transformers on the OBJECTID relationship "
            "with 100% match coverage — a meaningful improvement over "
            "the prior file-based approach.",
            "Added three identifier columns to downstream deliverables "
            "at the request of the field/GIS teams: SPID (service-point "
            "identifier), TransformerLID (transformer LID), and "
            "VaultCD. These now appear in every enriched output file so "
            "field crews and the application team can join model "
            "recommendations against their own systems more directly.",
            "Completed a memory-efficiency rewrite of the second "
            "(parallel) model's daily update path. Previously the "
            "weekly update needed a dedicated machine; now it runs "
            "alongside other work with roughly one-sixth the memory "
            "footprint. This unblocked two additional weekly runs "
            "and puts the parallel model on the same cadence as the "
            "production model.",
            "Grew the stability ledger from seven runs to nine. The "
            "additional history sharpens the high-confidence filter — "
            "recommendations that survive nine consecutive weeks of "
            "changing data are now treated as very strong evidence, "
            "while one-off flags are downweighted.",
            "Continued GIS data quality contributions. The model "
            "currently identifies 73 meters that are not yet in GIS "
            "and 1,290 meters whose GIS locations differ by more than "
            "100 meters from the source measurements (some by more "
            "than 20 kilometers, indicating likely data-entry errors).",
        ],
    )

    # ---- 6. Triage / Field Priorities ----
    add_heading(doc, "6. Triage Priority for Field & GIS Teams", level=1)
    add_para(
        doc,
        "If a field tech or GIS data steward asks which recommendations "
        "to investigate first, the priority order is:",
    )
    add_bullets(
        doc,
        [
            "Recommendations where both models agree at high confidence "
            "AND the model points across feeders — strongest possible "
            "signal that GIS has the feeder/transformer assignment "
            "wrong.",
            "Recommendations where both models agree at high confidence "
            "AND the meter is not yet in GIS — clean first-time "
            "mapping candidates for the GIS team.",
            "Recommendations where only the production model is "
            "high-confidence but the cross-check model sees the same "
            "case at lower confidence — strong but worth a closer look.",
            "Recommendations where only one model flags the case at "
            "high confidence — lower priority; treat as candidates "
            "rather than definite findings.",
            "Same-feeder ambiguous recommendations should never be "
            "actioned on model output alone. They reflect cases where "
            "voltage signatures cannot reliably distinguish adjacent "
            "transformers, and they require field validation.",
        ],
    )

    # ---- 7. What's Next ----
    add_heading(doc, "7. What's Next", level=1)
    add_bullets(
        doc,
        [
            "Continue the weekly run cadence for both models. Each "
            "additional run sharpens the stability scores and the "
            "consensus tiers; the parallel model is now on the same "
            "weekly cadence as production.",
            "Request the four missing decoded columns from the GIS "
            "dataflow owner (TRANSBANKTAG in ServicePoints; TAG, "
            "STRUCTNO, d_FEEDERID, d_SUBTYPECD in Transformers). The "
            "model works around their absence today using a fallback "
            "lookup from prior GIS extracts, but adding them to the "
            "dataflow will eliminate blanks for any new infrastructure "
            "created after 2026-04-28.",
            "Continue accumulating stability history. At six or more "
            "additional runs, the parallel model's own stability "
            "filter will begin producing meaningful signal, at which "
            "point the consensus tier becomes the most reliable "
            "output the model can produce.",
            "Maintain the structured feedback loop: as field "
            "operations and GIS work the recommendations, the "
            "suppression list and the field-validation observations "
            "feed back into future model runs.",
            "Package the weekly run sequence as a single command so "
            "the ordering (rolling window update → cluster rebuild → "
            "downstream reports → parallel model catch-up → consensus) "
            "is enforced by the tooling rather than by memory.",
        ],
    )

    add_para(doc, "")
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = closing.add_run(
        "Please reach out with any questions, requests for deeper "
        "drill-downs into specific recommendations, or feedback from "
        "field operations that should be captured in the suppression "
        "list or used to inform future model tuning."
    )
    r.italic = True
    r.font.size = Pt(11)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
